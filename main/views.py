import datetime
from django.shortcuts import render, redirect,get_object_or_404, Http404
from django.http import HttpResponse, JsonResponse,Http404
from .models import StudentBasicDetail, LeaveHistory, StudentContactDetail, FacultyDetail
from .serializers import LeaveHistorySerializer, StudentBasicDetailSerializer
from rest_framework.decorators import api_view
from rest_framework.response import Response
from django.contrib import messages
from datetime import date
from docx import Document
from django.core.mail import send_mail
from django.conf import settings
import json

# def student(request):
#     if request.method == 'GET':
#         print("get method hai")
#         return render(request, 'Student.html')
#     if request.method=="POST":
#         sp_id=request.session.get("sp_id")
#         print("hello",sp_id)
     
def hostel_incharge(request):
    if request.session.get('fp_id') is None:
        return redirect('login')
    if request.method == 'GET':
        today_leave = LeaveHistory.objects.filter(start_date__lte=date.today(), end_date__gte=date.today(),status_parent="approved", status_incharge="approved", status_fa="approved")
        today_leave_count = today_leave.count()
        context = {
            'today_leave': today_leave,
            'today_leave_count': today_leave_count
        }
        return render(request, 'Hostel.html', context)

def login(request):
    if request.method == 'GET':
        return render(request, 'login.html')

    elif request.method == 'POST':
        user_id = request.POST.get('sp_id')
        password = request.POST.get('password')

        try:
            hostel_incharge = get_object_or_404(FacultyDetail, fp_id=user_id)
            if hostel_incharge.password == password and hostel_incharge.department.dep_name == "Sharda Hostel Incharge":
                request.session['fp_id'] = user_id
                return redirect('hostel_incharge')
            else:
                messages.error(request, "Login failed. Incorrect password.")
                raise Http404
        except Http404:
            try:
                faculty = get_object_or_404(FacultyDetail, fp_id=user_id)
                if faculty.password == password:
                    request.session['fp_id'] = user_id
                    return redirect('faculty_advisor')
                else:
                    print("here")
                    messages.error(request, "Login failed. Incorrect password.")
                    return redirect('login')
            except Http404:
                try:
                    print("here")
                    student = get_object_or_404(StudentBasicDetail, sp_id=user_id)
                    print(student)
                    if student.password == password:
                        request.session['sp_id'] = user_id
                        print("here in student")
                        return redirect('student_login')
                        
                    else:
                        messages.error(request, "Login failed. Incorrect password.")
                        return redirect('login')
                except Http404:
                    messages.error(request, "Login failed. User not found.")
                    return redirect('login')


def logout(request):
    request.session.flush()
    return redirect('login')

def student_login(request):
    if request.session.get('sp_id') is None:
        return redirect('login')
    id = request.session.get('sp_id')
    print(id)
    student = get_object_or_404(StudentBasicDetail, sp_id=id)

    c_detail=StudentContactDetail.objects.get(StudentID=student.id)
    current_leave = LeaveHistory.objects.filter(student_id=student.id).last()
    leave_history = LeaveHistory.objects.filter(student_id=student.id, status_parent="approved", status_incharge="approved", status_fa="approved")
    return render(request,"Student.html",{"student":student,"info":c_detail,"leave":current_leave,"leave_history":leave_history})



def leaveapply(request):
    if request.session.get('sp_id') is None:
        return redirect('login')
    elif request.method == 'GET': #:
        return render(request, 'apply_leave.html')
    elif request.method == 'POST':
        sp_id = request.session.get('sp_id') #Get student id from session.
        start_date_str = request.POST.get('start_date')
        end_date_str = request.POST.get('end_date')
        reason = request.POST.get('reason')
        block = request.POST.get('block')
        room = request.POST.get('room')
        visiting_address = request.POST.get('visiting_address')

        try:
            start_date = datetime.datetime.strptime(start_date_str, '%Y-%m-%d').date()
            end_date = datetime.datetime.strptime(end_date_str, '%Y-%m-%d').date()

            no_of_days = (end_date - start_date).days + 1 #Add 1 to include the end date.

            if no_of_days <= 0:
                return HttpResponse("End date must be on or after start date.")

            student = StudentBasicDetail.objects.get(sp_id=sp_id)

            leave = LeaveHistory.objects.create(
                student_id=student,
                start_date=start_date,
                end_date=end_date,
                no_of_days=no_of_days,
                reason=reason,
                place_to_visit=visiting_address, 
            )

            if leave:
                details = StudentContactDetail.objects.get(StudentID=student)
                print(details.Parent_Email)
                message_content = {
                    'subject': 'Leave Application',
                    'body': f"Dear Parents, \n\n Your child {student.name} has applied for leave from {start_date} to {end_date}. \n\nReason: {reason}\n\nPlease approve or reject the leave application.\n\n Link: http://127.0.0.1:8000/parent_approve/{leave.id}\n\nBest regards,\n\nSharda Hostel",
                    'to': details.Parent_Email
                }
                send_email(message_content)

                return HttpResponse("Leave Application Successful")
            else:
                return HttpResponse("Leave Application Failed")

        except ValueError:
            return HttpResponse("Invalid date format. Please use YYYY-MM-DD.")
        except StudentBasicDetail.DoesNotExist:
            return HttpResponse("Student not found.")
        except Exception as e:
            return HttpResponse(f"An error occurred: {e}")


def faculty_advisor(request):
    if request.session.get('fp_id') is None:
        return redirect('login')
    id = request.session.get('fp_id')
    faculty_advisor = get_object_or_404(FacultyDetail, fp_id=id)
    if faculty_advisor.is_incharge == True:
        today = date.today()
        leave_applications = LeaveHistory.objects.filter(status_fa='approved', status_incharge='pending')
        pending_count = LeaveHistory.objects.filter(status_incharge='pending', status_fa='approved').count()
        approved_count = LeaveHistory.objects.filter(status_incharge='approved', status_fa='approved',start_date__gte=today).count()
        rejected_count = LeaveHistory.objects.filter(status_incharge='rejected', status_fa='approved').count()
        context = {
            'faculty_advisor': faculty_advisor,
            'leave_applications': leave_applications,
            'pending_count': pending_count,
            'approved_count': approved_count,
            'rejected_count': rejected_count,
        }

        return render(request, 'faculty_advisor.html', context)
    else:
        students = StudentBasicDetail.objects.filter(Batch_Year=faculty_advisor.Batch_Year, department=faculty_advisor.department)

    today = date.today()
    leave_applications = LeaveHistory.objects.filter(student_id__in=students,status_parent='approved', end_date__gte=today, status_fa='pending')

    pending_count = LeaveHistory.objects.filter(status_parent='approved', end_date__gte=today, status_fa='pending').count()
    approved_count = LeaveHistory.objects.filter(status_fa='approved',start_date__gte=today, end_date__gte=today).count()
    rejected_count = LeaveHistory.objects.filter(status_fa='rejected',start_date__gte=today, end_date__gte=today).count()

    context = {
        'faculty_advisor': faculty_advisor,
        'leave_applications': leave_applications,
        'pending_count': pending_count,
        'approved_count': approved_count,
        'rejected_count': rejected_count,
    }

    return render(request, 'faculty_advisor.html', context)


def approve_leave(request, leave_id):
    if request.method == 'POST':
        if request.session.get('fp_id') is None:
            return redirect('login')
        try:
            leave = get_object_or_404(LeaveHistory, id=leave_id)
            id = request.session.get('fp_id')
            faculty_advisor = get_object_or_404(FacultyDetail, fp_id=id)
            if faculty_advisor.is_incharge == True:
                leave.status_incharge = 'approved'
            else:
                leave.status_fa = 'approved'
            leave.save()
            return JsonResponse({'message': 'Leave approved successfully', 'ok': True}, status=200)
        except Exception as e:
            return JsonResponse({'error': str(e), 'ok': False}, status=500)
    return JsonResponse({'error': 'Invalid request method', 'ok': False}, status=400)

def reject_leave(request, leave_id):
    if request.method == 'POST':
        if request.session.get('fp_id') is None:
            return redirect('login')
        try:
            leave = get_object_or_404(LeaveHistory, id=leave_id)
            id = request.session.get('fp_id')
            faculty_advisor = get_object_or_404(FacultyDetail, fp_id=id)
            if faculty_advisor.is_incharge == True:
                leave.status_incharge = 'rejected'
            else:
                leave.status_fa = 'rejected'
            leave.save()
            return JsonResponse({'message': 'Leave rejected successfully', 'ok': True}, status=200)
        except Exception as e:
            return JsonResponse({'error': str(e), 'ok': False}, status=500)
    return JsonResponse({'error': 'Invalid request method', 'ok': False}, status=400)

def download_leave_letter(request, leave_id):
    if request.session.get('sp_id') is None:
        return redirect('login')
    else: 
        leave = get_object_or_404(LeaveHistory, id=leave_id)
        student = leave.student_id  # Assuming a ForeignKey relationship
        print(student.id)
        StudentContactInfo = get_object_or_404(StudentContactDetail, StudentID=student.id)
        document = Document()
        document.add_heading("Leave Letter", level=1)

        paragraph = document.add_paragraph()
        paragraph.add_run("Respected Madam, ").bold = True
        # paragraph.add_run(f"{student.name}").bold = True

        paragraph = document.add_paragraph()
        paragraph.add_run(f"I {student.name} studing in {student.department} request you to kindly accept my leave.").bold = True

        paragraph.add_run(f"My leave is from {leave.start_date} to {leave.end_date}, failing to which I will accept any academic penalty/or monetary penalty.").bold = True

        paragraph = document.add_paragraph()
        paragraph.add_run(f"Reason for seeking leave:{leave.reason}").bold = True

        paragraph = document.add_paragraph()
        paragraph.add_run(f"Hostel Block: {StudentContactInfo.Hostel_Block}").bold = True

        paragraph = document.add_paragraph()
        paragraph.add_run(f"Room Number: {StudentContactInfo.Room_No}, Hostel Name: Sharda Girls Hostel").bold = True

        paragraph = document.add_paragraph()
        paragraph.add_run(f"Student Contact Number: {StudentContactInfo.PhoneNumber}").bold = True

        paragraph = document.add_paragraph()
        paragraph.add_run(f"Parent Contact Number: {StudentContactInfo.Parent_No}").bold = True

        paragraph = document.add_paragraph()
        paragraph.add_run("Thank you for your understanding.").bold = True

        paragraph = document.add_paragraph()
        paragraph.add_run("Sincerely,").bold = True

        paragraph = document.add_paragraph()
        paragraph.add_run(f"{student.name}").bold = True

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=leave_letter.docx'
        document.save(response)
        return response


# Email function
def send_email(message_content):
    subject = message_content['subject']
    message = message_content['body']
    email_to = message_content['to']
    email_from = settings.EMAIL_HOST_USER
    

    try:
        send_mail(subject, message, email_from, [email_to], fail_silently=False)
        return HttpResponse("Email sent successfully")
    except Exception as e:
        return HttpResponse(f"Error sending email: {e}")


def parent_approve(request,leave_id):
    if request.method == 'GET':
        leave = LeaveHistory.objects.filter(id=leave_id,status_parent="pending")
        print(leave)
        context = {
            'leave':leave
        }
        return render(request, 'parent_approve.html', context)
    elif request.method == 'POST':
        body = json.loads(request.body)
        print(body['action'])
        if body['action'] == 'approve':
            leave = get_object_or_404(LeaveHistory, id=leave_id)
            leave.status_parent = 'approved'
            leave.save()
        elif body['action'] == 'reject':
            leave = get_object_or_404(LeaveHistory, id=leave_id)
            leave.status_parent = 'rejected'
            leave.save()
        return JsonResponse({'message': 'Leave rejected successfully', 'ok': True}, status=200)